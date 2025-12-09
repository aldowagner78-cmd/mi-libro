#!/usr/bin/env python3
"""
Analiza todos los archivos DOCX y reporta problemas de formato.
"""

import mammoth
import re
from pathlib import Path
from docx import Document

def analyze_docx(docx_path):
    """Analiza un DOCX y reporta su estructura."""
    
    chapter_name = Path(docx_path).stem
    print(f"\n{'='*70}")
    print(f"📄 ANALIZANDO: {chapter_name}")
    print('='*70)
    
    # Convertir con Mammoth para ver HTML
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)
        html = result.value
    
    # 1. Buscar referencias SUP en el texto
    sups = re.findall(r'<sup>([^<]+)</sup>', html)
    refs_con_parentesis = [s for s in sups if re.match(r'\(\d+\)', s.strip())]
    
    print(f"\n📌 REFERENCIAS EN EL TEXTO:")
    print(f"   Total <sup> encontrados: {len(sups)}")
    if refs_con_parentesis:
        print(f"   ✅ Con formato (N): {len(refs_con_parentesis)}")
        print(f"      Números: {', '.join(refs_con_parentesis[:10])}")
        if len(refs_con_parentesis) > 10:
            print(f"      ... y {len(refs_con_parentesis) - 10} más")
    else:
        print(f"   ❌ Sin referencias con formato (N)")
    
    # Referencias sin formato correcto
    refs_malas = [s for s in sups if not re.match(r'\(\d+\)', s.strip())]
    if refs_malas:
        print(f"   ⚠️  Referencias con formato incorrecto: {refs_malas[:5]}")
    
    # 2. Buscar sección de notas
    tiene_seccion_notas = "NOTAS DEL" in html.upper()
    print(f"\n📝 SECCIÓN DE NOTAS:")
    
    if tiene_seccion_notas:
        print(f"   ✅ Encontrada sección 'NOTAS DEL...'")
        
        # Extraer notas
        if "NOTAS DEL CAPÍTULO" in html or "NOTAS DEL" in html:
            parts = re.split(r'<p><strong>NOTAS DEL[^<]*</strong></p>', html, maxsplit=1)
            
            if len(parts) == 2:
                notes_html = parts[1]
                
                # Formato 1: <p><sup>(N)</sup> ...
                pattern1 = r'<p><sup>\((\d+)\)</sup>'
                notas_formato1 = re.findall(pattern1, notes_html)
                
                # Formato 2: <p><strong>(N) ...
                pattern2 = r'<p><strong>\((\d+)\)'
                notas_formato2 = re.findall(pattern2, notes_html)
                
                total_notas = len(set(notas_formato1 + notas_formato2))
                
                if total_notas > 0:
                    print(f"   ✅ Notas detectadas: {total_notas}")
                    
                    if notas_formato1:
                        print(f"      Formato <sup>(N)</sup>: {len(notas_formato1)}")
                        print(f"         Números: {', '.join(notas_formato1)}")
                    
                    if notas_formato2:
                        print(f"      Formato <strong>(N)</strong>: {len(notas_formato2)}")
                        print(f"         Números: {', '.join(notas_formato2)}")
                    
                    # Verificar que coincidan
                    nums_refs = set([s.strip('()') for s in refs_con_parentesis])
                    nums_notas = set(notas_formato1 + notas_formato2)
                    
                    faltantes = nums_refs - nums_notas
                    sobrantes = nums_notas - nums_refs
                    
                    if faltantes:
                        print(f"   ⚠️  Referencias sin nota: {', '.join(sorted(faltantes, key=int))}")
                    
                    if sobrantes:
                        print(f"   ⚠️  Notas sin referencia: {', '.join(sorted(sobrantes, key=int))}")
                    
                    if not faltantes and not sobrantes:
                        print(f"   ✅ Todas las referencias tienen su nota")
                else:
                    print(f"   ❌ No se detectaron notas (revisar formato)")
            else:
                print(f"   ❌ No se pudo separar la sección de notas")
    else:
        print(f"   ❌ No se encontró sección 'NOTAS DEL...'")
        if len(refs_con_parentesis) > 0:
            print(f"   ⚠️  HAY REFERENCIAS PERO NO HAY SECCIÓN DE NOTAS")
    
    # 3. Revisar estructura con python-docx
    try:
        doc = Document(docx_path)
        
        print(f"\n📐 ESTRUCTURA DEL DOCUMENTO:")
        print(f"   Párrafos totales: {len(doc.paragraphs)}")
        
        # Buscar estilos usados
        estilos = set([p.style.name for p in doc.paragraphs if p.text.strip()])
        print(f"   Estilos usados: {', '.join(sorted(estilos)[:5])}")
        
        # Buscar notas automáticas de Word
        if doc.footnotes:
            print(f"   ⚠️  NOTAS AUTOMÁTICAS DE WORD DETECTADAS: {len(doc.footnotes)}")
            print(f"      ELIMÍNALAS: no son compatibles con el lector web")
        
    except Exception as e:
        print(f"   ⚠️  No se pudo analizar estructura: {e}")
    
    # 4. Diagnóstico final
    print(f"\n🎯 DIAGNÓSTICO:")
    
    if tiene_seccion_notas and total_notas > 0:
        if not faltantes and not sobrantes:
            print(f"   ✅ FORMATO CORRECTO - Todo funcionará")
        else:
            print(f"   ⚠️  REVISAR - Hay inconsistencias entre referencias y notas")
    elif tiene_seccion_notas and total_notas == 0:
        print(f"   ❌ FORMATO INCORRECTO - Sección de notas no se puede leer")
        print(f"      Verifica que las notas empiecen con (1), (2), etc.")
    elif not tiene_seccion_notas and len(refs_con_parentesis) > 0:
        print(f"   ❌ FALTA SECCIÓN DE NOTAS")
        print(f"      Agrega 'NOTAS DEL CAPÍTULO X' al final del documento")
    elif not tiene_seccion_notas and len(refs_con_parentesis) == 0:
        print(f"   ℹ️  Sin notas (esto está bien si no hay referencias)")
    else:
        print(f"   ⚠️  Revisar manualmente el documento")

def main():
    print("="*70)
    print("🔍 ANÁLISIS DE FORMATO DE ARCHIVOS DOCX")
    print("="*70)
    
    capitulos_dir = Path("capitulos")
    docx_files = sorted(capitulos_dir.glob("*.docx"))
    
    for docx_file in docx_files:
        try:
            analyze_docx(docx_file)
        except Exception as e:
            print(f"\n❌ ERROR: {e}")
    
    print("\n" + "="*70)
    print("✅ ANÁLISIS COMPLETADO")
    print("="*70)
    print("\n📖 Lee GUIA_FORMATO_DOCX.md para ver cómo corregir los problemas")

if __name__ == "__main__":
    main()
